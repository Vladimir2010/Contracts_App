import os
from datetime import datetime
from docx import Document
from typing import Dict, Any, List
import locale

# Try to set Bulgarian locale for dates
try:
    locale.setlocale(locale.LC_ALL, 'bg_BG.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_ALL, 'bulgarian')
    except:
        pass

def format_date_long_bg(dt=None):
    """Format date as: понеделник, 12.януари. 2026 г."""
    if dt is None:
        dt = datetime.now()
    
    days = ["понеделник", "вторник", "сряда", "четвъртък", "петък", "събота", "неделя"]
    months = ["януари", "февруари", "март", "април", "май", "юни", 
              "юли", "август", "септември", "октомври", "ноември", "декември"]
    
    day_name = days[dt.weekday()]
    month_name = months[dt.month - 1]
    
    return f"{day_name}, {dt.day}.{month_name}. {dt.year} г."

import re

def clean_xml_string(s):
    """Aggressively remove any character that could break Word's XML 1.0 body"""
    if s is None: return ""
    s = str(s)
    # XML 1.0 valid chars: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    # We use a whitelist approach for maximum safety.
    return "".join(c for c in s if ord(c) in (0x9, 0xA, 0xD) or (0x20 <= ord(c) <= 0xD7FF) or (0xE000 <= ord(c) <= 0xFFFD))

def surgical_replace(para, placeholder, replacement, once=False):
    """
    The most robust way to replace text in python-docx:
    1. If placeholder is in a single run, replace it there.
    2. If split across runs, merge into the first run and clear others.
    Preserves all Run objects to avoid XML handle corruption.
    """
    if placeholder not in para.text:
        return False
        
    replacement = clean_xml_string(replacement)
    
    # Try one-run replacement first
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement, 1 if once else -1)
            return True
            
    # Handle split runs (the 'Safe Merge' strategy)
    full_text = para.text
    if once:
        new_text = full_text.replace(placeholder, replacement, 1)
    else:
        new_text = full_text.replace(placeholder, replacement)
        
    if new_text != full_text and para.runs:
        # Move all text to run 0, clear others. 
        # This keeps the XML structure identical but changes the content.
        para.runs[0].text = new_text
        for i in range(1, len(para.runs)):
            para.runs[i].text = ""
        return True
        
    return False

def replace_text_once(doc, placeholder, replacement):
    """Surgical replacement of the first occurrence found in the document"""
    for para in doc.paragraphs:
        if surgical_replace(para, placeholder, replacement, once=True):
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if surgical_replace(para, placeholder, replacement, once=True):
                        return True
    return False

def replace_text_all(doc, placeholder, replacement):
    """Surgical replacement of all occurrences in the document"""
    found = False
    for para in doc.paragraphs:
        if surgical_replace(para, placeholder, replacement, once=False):
            found = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if surgical_replace(para, placeholder, replacement, once=False):
                        found = True
    return found

def doc_to_docx(doc_path):
    """Convert .doc to .docx using pywin32 with robust cleanup"""
    import win32com.client
    import pythoncom
    
    # Initialize COM for the thread
    pythoncom.CoInitialize()
    
    word = None
    doc = None
    try:
        # Use DispatchEx to ensure a fresh instance
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0 # wdAlertsNone
        
        doc_path = os.path.abspath(doc_path)
        docx_path = doc_path + "x"
        
        print(f"Конвертиране на: {doc_path}")
        
        # Open the document (ReadOnly to avoid lock issues)
        doc = word.Documents.Open(doc_path, ReadOnly=True, ConfirmConversions=False)
        
        # Save as docx (FileFormat 16)
        doc.SaveAs2(docx_path, FileFormat=16)
        
        doc.Close(False)
        doc = None
        
        return docx_path
    except Exception as e:
        print(f"Грешка при COM конвертиране: {e}")
        raise e
    finally:
        try:
            if doc: doc.Close(False)
        except: pass
        try:
            if word: word.Quit()
        except: pass
        pythoncom.CoUninitialize()

def clean_numeric(val):
    """Remove .0 from numbers like serials or fiscal memory"""
    if val is None: return ""
    s = str(val).strip()
    if s.endswith('.0'):
        return s[:-2]
    return s

def format_phone_custom(phone):
    """Format phone as 0888/728-005 or 02/870-5657"""
    if not phone: return ""
    digits = re.sub(r'\D', '', str(phone))
    if not digits: return str(phone)
    
    # 0888728005 -> 0888/728-005
    if len(digits) == 10 and digits.startswith('08'):
        return f"{digits[:4]}/{digits[4:7]}-{digits[7:]}"
    # 028705657 -> 02/870-5657
    if len(digits) == 9 and digits.startswith('02'):
        return f"{digits[:2]}/{digits[2:5]}-{digits[5:]}"
    # Fallback for other formats
    if len(digits) > 6:
        return f"{digits[:-6]}/{digits[-6:-3]}-{digits[-3:]}"
    return str(phone)

def format_date_bg(dt, fmt_type='A'):
    """
    fmt_type A: 15/01/26 г.
    fmt_type B: 15 януари 2026 г.
    fmt_type C: четвъртък, 15 януари 2026 г.
    fmt_type D: 15.01.2027 г. (for device list)
    """
    if not dt or not isinstance(dt, datetime):
        return ""
    
    days = ["понеделник", "вторник", "сряда", "четвъртък", "петък", "събота", "неделя"]
    months = ["януари", "февруари", "март", "април", "май", "юни", 
              "юли", "август", "септември", "октомври", "ноември", "декември"]
    
    if fmt_type == 'A':
        return dt.strftime('%d/%m/%y г.')
    elif fmt_type == 'B':
        return f"{dt.day} {months[dt.month - 1]} {dt.year} г."
    elif fmt_type == 'C':
        return f"{days[dt.weekday()]}, {dt.day} {months[dt.month - 1]} {dt.year} г."
    elif fmt_type == 'D':
        return dt.strftime('%d.%m.%Y г.')
    return dt.strftime('%d.%m.%Y')

def docx_to_pdf(docx_path):
    """Convert .docx to .pdf using pywin32"""
    import win32com.client
    import pythoncom
    
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        
        docx_path = os.path.abspath(docx_path)
        pdf_path = docx_path.rsplit('.', 1)[0] + ".pdf"
        
        doc = word.Documents.Open(docx_path, ReadOnly=True)
        # wdFormatPDF = 17
        doc.SaveAs2(pdf_path, FileFormat=17)
        doc.Close(False)
        return pdf_path
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        return None
    finally:
        if word: word.Quit()
        pythoncom.CoUninitialize()

def number_to_words_bg(amount, currency="BGN"):
    """
    Convert number to Bulgarian words for currency amounts.
    Handle BGN (лева/стотинки) and EUR (евро/цента).
    """
    if amount is None or amount == "": return ""
    try:
        amount = float(amount)
    except:
        return str(amount)

    units = ["", "един", "два", "три", "четири", "пет", "шест", "седем", "осем", "девет"]
    units_fem = ["", "една", "две", "три", "четири", "пет", "шест", "седем", "осем", "девет"]
    teens = ["десет", "единадесет", "дванадесет", "тринадесет", "четиринадесет", "петнадесет", "шестнадесет", "седемнадесет", "осемнадесет", "деветнадесет"]
    tens = ["", "десет", "двадесет", "тридесет", "четиридесет", "петдесет", "шестдесет", "седемдесет", "осемдесет", "деветдесет"]
    hundreds = ["", "сто", "двеста", "триста", "четиристотин", "петстотин", "шестстотин", "седемстотин", "осемстотин", "деветстотин"]

    def convert_chunk(num, gender='masc'):
        res = []
        h = num // 100
        t = (num % 100) // 10
        u = num % 10
        
        if h > 0: res.append(hundreds[h])
        
        target_units = units_fem if gender == 'fem' else units
        
        if t == 1:
            if h > 0: res.append("и")
            res.append(teens[u])
        else:
            if t > 0:
                if h > 0: res.append("и")
                res.append(tens[t])
                if u > 0:
                    res.append("и")
                    res.append(target_units[u])
            elif u > 0:
                if h > 0: res.append("и")
                res.append(target_units[u])
        return " ".join(res)

    integer_part = int(amount)
    fraction_part = round((amount - integer_part) * 100)

    # Simplified Bulgarian word conversion for amounts
    parts = []
    
    # Millions
    mil = integer_part // 1000000
    if mil > 0:
        if mil == 1: parts.append("един милион")
        else: parts.append(convert_chunk(mil, 'masc') + " милиона")
    
    # Thousands
    thousands = (integer_part % 1000000) // 1000
    if thousands > 0:
        if thousands == 1: parts.append("хиляда")
        else: parts.append(convert_chunk(thousands, 'fem') + " хиляди")
        
    # Basics
    rest = integer_part % 1000
    if rest > 0 or not parts:
        if parts and rest < 100: parts.append("и")
        parts.append(convert_chunk(rest, 'masc'))

    words = " ".join(parts).strip()
    
    if currency == "BGN":
        main_unit = "лев" if integer_part == 1 else "лева"
        frac_unit = "стотинка" if fraction_part == 1 else "стотинки"
        return f"{words} {main_unit} и {fraction_part:02d} {frac_unit}"
    else: # EUR
        main_unit = "евро"
        frac_unit = "цента"
        return f"{words} {main_unit} и {fraction_part:02d} {frac_unit}"

def generate_service_contract(client_data: Dict[str, Any], devices: List[Dict[str, Any]], template_path: str, output_dir: str) -> str:
    """
    Generate service contract using strict {1}-{51} placeholder mapping.
    """
    from path_utils import get_resource_path
    template_path = get_resource_path(template_path)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template НЕ е намерен: {template_path}")

    doc = Document(template_path)
    
    now = datetime.now()
    # Contract start date as datetime object for comparison
    start_date_str = client_data.get('contract_start', '')
    try:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
    except:
        start_date = now

    # Basic data
    c_num = str(client_data.get('contract_number', ''))
    c_name = str(client_data.get('company_name', ''))
    c_addr = str(client_data.get('address', ''))
    mol = str(client_data.get('mol', ''))
    eik_pure = clean_numeric(client_data.get('eik', ''))
    is_vat = str(client_data.get('vat_registered', '')).lower() == 'да'
    eik_val = f"BG{eik_pure}" if is_vat else eik_pure
    
    mappings = {
        "{1}": c_num,
        "{2}": format_date_bg(now, 'A'),
        "{3}": c_name,
        "{4}": c_addr,
        "{5}": eik_val,
        "{6}": format_phone_custom(client_data.get('phone1', '')),
        "{7}": mol,
        "{8}": format_date_bg(now, 'B'),
        "{9}": c_num,
        "{10}": format_date_bg(now, 'A'),
        "{46}": format_date_bg(now, 'B'),
        "{47}": c_num,
        "{48}": format_date_bg(now, 'A'),
        "{49}": format_date_bg(now, 'C'),
        "{50}": "Г" if start_date.date() == now.date() else "А"
    }

    # Device Mapping (11-45, grouped by 7 fields per device)
    for i in range(5):
        base_idx = 11 + (i * 7)
        if i < len(devices):
            dev = devices[i]
            mappings[f"{{{base_idx}}}"] = str(dev.get('object_name', ''))
            mappings[f"{{{base_idx+1}}}"] = str(dev.get('object_address', ''))
            mappings[f"{{{base_idx+2}}}"] = format_phone_custom(dev.get('object_phone', ''))
            mappings[f"{{{base_idx+3}}}"] = mol
            mappings[f"{{{base_idx+4}}}"] = str(dev.get('model', ''))
            mappings[f"{{{base_idx+5}}}"] = clean_numeric(dev.get('serial_number', ''))
            mappings[f"{{{base_idx+6}}}"] = clean_numeric(dev.get('fiscal_memory', ''))
        else:
            for j in range(7):
                mappings[f"{{{base_idx+j}}}"] = ""

    # {51} - Device list with expiry dates
    device_list_entries = []
    for i, dev in enumerate(devices):
        expiry_str = str(dev.get('contract_expiry', ''))
        expiry_formatted = ""
        try:
            exp_date = datetime.strptime(expiry_str, '%Y-%m-%d')
            expiry_formatted = format_date_bg(exp_date, 'D')
        except:
            expiry_formatted = expiry_str
        
        device_list_entries.append(f"ЕКА No {i+1} до {expiry_formatted}")
    
    mappings["{51}"] = ", ".join(device_list_entries)

    # Perform replacements
    for ph, val in mappings.items():
        replace_text_all(doc, ph, clean_xml_string(val))

    # Save
    safe_company = "".join([c for c in c_name if c.isalnum() or c in (' ', '-', '_')]).strip()
    output_filename = f"{c_num} {safe_company}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    doc.save(output_path)
    return output_path

def generate_registration_certificate(client_data, device, template_path, output_dir):
    """Generate RegCert_SN.docx from template"""
    from path_utils import get_resource_path
    template_path = get_resource_path(template_path)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template НЕ е намерен: {template_path}")

    doc = Document(template_path)
    now = datetime.now()
    
    sn = clean_numeric(device.get('serial_number', ''))
    
    # 1 - 15/01/2025г.
    # 12 - 15/01/2025 г.
    date_f1 = now.strftime('%d/%m/%Yг.')
    date_f12 = now.strftime('%d/%m/%Y г.')
    
    c_start = client_data.get('contract_start', '')
    try:
        dt_start = datetime.strptime(c_start, '%Y-%m-%d')
        start_fmt = dt_start.strftime('%d.%m.%Y г.')
    except:
        start_fmt = str(c_start)

    mappings = {
        "{1}": date_f1,
        "{2}": clean_numeric(client_data.get('eik', '')),
        "{3}": str(client_data.get('company_name', '')),
        "{4}": str(client_data.get('address', '')),
        "{5}": str(client_data.get('mol', '')),
        "{6}": f"{device.get('object_name', '')}, {device.get('object_address', '')}",
        "{7}": str(device.get('model', '')),
        "{8}": str(device.get('bim_number', '')),
        "{9}": sn,
        "{10}": clean_numeric(device.get('fiscal_memory', '')),
        "{11}": str(client_data.get('contract_number', '')),
        "{12}": date_f12,
        "{13}": clean_numeric(device.get('fdrid', '')),
        "{14}": start_fmt
    }
    
    for ph, val in mappings.items():
        replace_text_all(doc, ph, clean_xml_string(val))
        
    output_filename = f"RegCert_{sn}.docx"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    return output_path

def generate_deregistration_protocol(proto_data, template_path, output_dir):
    """Generate DeregProtocol_SN.docx from complex proto_data dict"""
    from path_utils import get_resource_path
    template_path = get_resource_path(template_path)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template НЕ е намерен: {template_path}")

    doc = Document(template_path)
    
    # Manufacturers Data
    # DY -> Daisy ("Дейзи Тех" АД ЕИК - 201679556)
    # DT -> Datecs ( "Датекс" ООД ЕИК - 000713391)
    # ZK / TR -> Tremol ("Тремол" ООД ЕИК - 104593442)
    sn = str(proto_data.get('serial_number', ''))
    manu_sel = proto_data.get('manufacturer', '')
    manu_name = ""
    manu_eik = ""
    manu_city = ""
    
    if manu_sel == "Дейзи" or (not manu_sel and (sn.startswith('DY') or sn.startswith('SY'))):
        manu_name = '"Дейзи Тех" АД'
        manu_eik = "201679556"
        manu_city = "София"
    elif manu_sel == "Датекс" or (not manu_sel and sn.startswith('DT')):
        manu_name = '"Датекс" ООД'
        manu_eik = "000713391"
        manu_city = "София"
    elif manu_sel == "Тремол" or (not manu_sel and (sn.startswith('ZK') or sn.startswith('TR') or sn.startswith('TE'))):
        manu_name = '"Тремол" ООД'
        manu_eik = "104593442"
        manu_city = "Велико Търново"
        
    now = datetime.now()
    curr_label = "лв." if proto_data.get('currency', 'BGN') == 'BGN' else "€"
    
    def fmt_amt(val):
        if not val: return f"0.00 {curr_label}"
        try:
            num = float(val)
            if curr_label == "€": return f"€ {num:.2f}"
            return f"{num:.2f} лв."
        except: return str(val)

    # Certificate Expiry Date
    cert_date = proto_data.get('certificate_expiry', None)
    bim_no = str(proto_data.get('bim_number', ''))
    
    if not cert_date and bim_no:
        try:
            from database import get_certificate_expiry
            cert_date = get_certificate_expiry(bim_no)
        except: pass
        
    date_f8 = ""
    if cert_date:
        try:
            if isinstance(cert_date, str) and '-' in cert_date:
                dt_cert = datetime.strptime(cert_date, '%Y-%m-%d')
            elif isinstance(cert_date, str) and '.' in cert_date:
                dt_cert = datetime.strptime(cert_date, '%d.%m.%Y')
            else:
                dt_cert = cert_date # Already a date object?
            date_f8 = dt_cert.strftime('%d.%m.%Y г.')
        except:
            date_f8 = str(cert_date)
    else:
        date_f8 = now.strftime('%d.%m.%Y г.') # Fallback
        
    mappings = {
        "{1}": now.strftime('%d.%m.%Y г.'),
        "{2}": now.strftime('%H:%M'),
        "{3}": str(proto_data.get('eik', '')),
        "{4}": f"{proto_data.get('company_name', '')}, {proto_data.get('address', '')}",
        "{5}": f"{proto_data.get('mol', '')}, {proto_data.get('address', '')}",
        "{6}": f"{proto_data.get('object_name', '')}, {proto_data.get('object_address', '')}",
        "{7}": str(proto_data.get('model', '')),
        "{8}": f"{bim_no} / {date_f8}",
        "{9}": manu_name,
        "{10}": manu_eik,
        "{11}": sn,
        "{12}": clean_numeric(proto_data.get('fiscal_memory', '')),
        "{13}": clean_numeric(proto_data.get('fdrid', '')),
        "{14}": str(proto_data.get('reason', '')),
        "{15}": proto_data.get('date_start_fmt', ''), 
        "{16}": proto_data.get('date_stop_fmt', ''),
        "{17}": fmt_amt(proto_data.get('turnover', 0)),
        "{18}": number_to_words_bg(proto_data.get('turnover', 0), proto_data.get('currency', 'BGN')),
        "{19}": fmt_amt(proto_data.get('turnover', 0)),
        "{20}": fmt_amt(proto_data.get('storno_total', 0)),
        "{21}": fmt_amt(proto_data.get('vat_a', 0)),
        "{22}": fmt_amt(proto_data.get('vat_b', 0)),
        "{23}": fmt_amt(proto_data.get('vat_v', 0)),
        "{24}": fmt_amt(proto_data.get('vat_g', 0)),
        "{25}": fmt_amt(proto_data.get('storno_a', 0)),
        "{26}": fmt_amt(proto_data.get('storno_b', 0)),
        "{27}": fmt_amt(proto_data.get('storno_v', 0)),
        "{28}": fmt_amt(proto_data.get('storno_g', 0)),
        "{29}": f"{manu_name}, гр. {manu_city}",
        "{30}": f"{proto_data.get('company_name', '')}, гр. София" # Defaulting to Sofia or client city
    }
    
    for ph, val in mappings.items():
        replace_text_all(doc, ph, clean_xml_string(val))
        
    output_filename = f"DeregProtocol_{sn}.docx"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    return output_path

def generate_nap_xml(service_data, client_eik, fdrid, output_dir):
    """Generate NAP XML file in WINDOWS-1251 encoding"""
    now = datetime.now()
    # NAP_YYYYMMDD_HHMMSS.xml
    timestamp = now.strftime('%Y%m%d_%H%M%S')
    filename = f"NAP_{timestamp}.xml"
    output_path = os.path.join(output_dir, filename)
    
    # Service Name formatting: Uppercase and no quotes
    service_name = str(service_data.get('name', '')).replace('"', '').replace("'", "").upper().strip()
    
    # Structure from example
    xml_content = f"""<?xml version="1.0" encoding="WINDOWS-1251"?>
<dec44a2 xmlns="http://inetdec.nra.bg/xsd/dec_44a2.xsd" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" schemaLocation="http://inetdec.nra.bg/xsd/dec_44a2.xsd http://inetdec.nra.bg/xsd/dec_44a2.xsd">
  <name>{service_name}</name>
  <bulstat>{service_data.get('eik', '')}</bulstat>
  <telcode>{service_data.get('phone1', '')}</telcode>
  <telnum>{service_data.get('phone2', '')}</telnum>
  <authorizeid>{service_data.get('tech_egn', '')}</authorizeid>
  <autorizecode>1</autorizecode>
  <fname>{clean_xml_string(service_data.get('tech_f', ''))}</fname>
  <sname>{clean_xml_string(service_data.get('tech_m', ''))}</sname>
  <tname>{clean_xml_string(service_data.get('tech_l', ''))}</tname>
  <id>{client_eik}</id>
  <code>5</code>
  <fuiasutd>
    <rowenum>
      <fdrid>{fdrid}</fdrid>
    </rowenum>
  </fuiasutd>
</dec44a2>"""

    try:
        with open(output_path, 'wb') as f:
            f.write(xml_content.encode('windows-1251', errors='replace'))
        return output_path
    except Exception as e:
        raise Exception(f"Грешка при запис на XML: {e}")
