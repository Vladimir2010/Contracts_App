from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_run_language_bg(run):
    """Force the run language to Bulgarian so Word evaluates fields in BG"""
    rPr = run._r.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'bg-BG')
    rPr.append(lang)

def add_auto_updating_date(paragraph, date_format="dd.MM.yyyy", use_switch=False):
    run = paragraph.add_run()
    set_run_language_bg(run)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    
    # \l 1026 forces Bulgarian locale for the field
    if use_switch:
        instrText.text = f' DATE \\@ "{date_format}" \\l 1026 \\* MERGEFORMAT '
    else:
        instrText.text = f' DATE \\@ "{date_format}" \\* MERGEFORMAT '
        
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    
    text_run = paragraph.add_run("ДД.ММ.ГГГГ")
    set_run_language_bg(text_run)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    text_run._r.append(fldChar3)

def main():
    doc = Document()
    doc.add_heading('Тест - Дати на Български', level=1)
    
    doc.add_paragraph('За да се показват на български, задаваме езика на текста (bg-BG) и евентуално използваме код \\l 1026 в полето.')
    
    doc.add_paragraph('=== 1. Само със зададен език на текста (bg-BG) ===')
    p2 = doc.add_paragraph('Формат B (d MMMM yyyy г.): ')
    add_auto_updating_date(p2, "d MMMM yyyy 'г.'", use_switch=False)
    
    p3 = doc.add_paragraph('Формат C (dddd, d MMMM yyyy г.): ')
    add_auto_updating_date(p3, "dddd, d MMMM yyyy 'г.'", use_switch=False)
    
    doc.add_paragraph('=== 2. Със зададен език + код \\l 1026 ===')
    p4 = doc.add_paragraph('Формат B (d MMMM yyyy г.): ')
    add_auto_updating_date(p4, "d MMMM yyyy 'г.'", use_switch=True)
    
    p5 = doc.add_paragraph('Формат C (dddd, d MMMM yyyy г.): ')
    add_auto_updating_date(p5, "dddd, d MMMM yyyy 'г.'", use_switch=True)

    doc.add_paragraph('=== 3. Останалите формати (С код \\l 1026) ===')
    p6 = doc.add_paragraph('Формат A (dd/MM/yy г.): ')
    add_auto_updating_date(p6, "dd/MM/yy 'г.'", use_switch=True)
    
    p7 = doc.add_paragraph('Формат D (dd.MM.yyyy г.): ')
    add_auto_updating_date(p7, "dd.MM.yyyy 'г.'", use_switch=True)
    
    doc.save('test_auto_date3.docx')
    print("Документът е запазен като test_auto_date3.docx")

if __name__ == "__main__":
    main()
