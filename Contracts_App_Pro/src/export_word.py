from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Tuple


def export_to_word(data: List[Tuple], headers: List[str], filename: str, title: str = "Справка за изтичащи договори") -> bool:
    """
    Export data to Word document with formatted table.
    
    Args:
        data: List of tuples containing row data
        headers: List of column headers
        filename: Output filename (should end with .docx)
        title: Document title
    
    Returns:
        True if successful, False otherwise
    """
    try:
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add empty line
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = header
            # Make header bold
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row_data):
                row_cells[idx].text = str(value) if value else ""
        
        # Save document
        doc.save(filename)
        return True
    
    except Exception as e:
        print(f"Error exporting to Word: {e}")
        return False
