import os
from docx import Document
from datetime import datetime
from path_utils import get_app_root

def generate_handover_protocol(data: dict, output_filename: str = None):
    """
    Generate a Word document based on the template in resources/
    Replaces markers {1} to {15} with data.
    """
    template_path = os.path.join(get_app_root(), "resources", "priemo_predavatelen_protokol.docx")
    
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return None
        
    try:
        doc = Document(template_path)
        
        # Prepare replacements mapping
        replacements = {
            "{1}": data.get('date', datetime.now().strftime("%d.%m.%Y")),
            "{2}": "гр. София",
            "{3}": data.get('service_firm', ''),
            "{4}": data.get('service_eik', ''),
            "{5}": data.get('service_address', ''),
            "{6}": data.get('service_mol', ''),
            "{7}": data.get('tech_egn', ''),
            "{8}": data.get('capacity', 'Сервизен техник'),
            "{9}": data.get('client_name', ''),
            "{10}": data.get('client_eik', ''),
            "{11}": data.get('client_address', ''),
            "{12}": data.get('client_mol', ''),
            "{13}": data.get('description', ''),
            "{14}": data.get('notes', ''),
            "{15}": data.get('ref_number', '')
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, val in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(val))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, val in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, str(val))
        
        # Save output
        if not output_filename:
            save_dir = os.path.join(os.path.expanduser("~"), "Documents", "ContractsApp", "Protocols")
            os.makedirs(save_dir, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = os.path.join(save_dir, f"Protocol_{timestamp}.docx")
            
        doc.save(output_filename)
        return output_filename
    except Exception as e:
        print(f"Error generating protocol: {e}")
        return None
