import os
from docx import Document

def doc_to_docx(doc_path):
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_path = os.path.abspath(doc_path)
    docx_path = doc_path + "x"
    if os.path.exists(docx_path): os.remove(docx_path)
    doc = word.Documents.Open(doc_path)
    doc.SaveAs2(docx_path, FileFormat=16) 
    doc.Close()
    word.Quit()
    return docx_path

def inspect_runs(path):
    if path.lower().endswith('.doc'):
        path = doc_to_docx(path)
    doc = Document(path)
    print("--- TARGETED SEARCH ---")
    for i, para in enumerate(doc.paragraphs):
        t = para.text
        if any(kw in t for kw in ["дванадесет", "ЗДДС", "BG", "ЕИК"]):
            print(f"{i}: {t}")

    print("\n--- SEARCHING ALL TABLES ---")
    for i, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                if any(kw in cell.text for kw in ["ЕИК", "ЗДДС", "БУЛСТАТ"]):
                    print(f"T{i} R{r} C{c}: {cell.text}")

    print("\n--- TABLES (FIRST ROW) ---")
    for i, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows[:2]):
            for c, cell in enumerate(row.cells):
                if any(ch in cell.text for ch in "()"):
                    print(f"T{i} R{r} C{c}: {cell.text}")

if __name__ == "__main__":
    inspect_runs("1 Профинанс Д и Д ЕООД.doc")
