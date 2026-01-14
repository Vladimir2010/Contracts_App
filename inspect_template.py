import os
from docx import Document

def doc_to_docx(doc_path):
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_path = os.path.abspath(doc_path)
    docx_path = doc_path + "x"
    if os.path.exists(docx_path): os.remove(docx_path)
    doc = word.Documents.Open(doc_path)
    doc.SaveAs2(docx_path, FileFormat=16) 
    doc.Close()
    word.Quit()
    return docx_path

def inspect_docx(path):
    if path.lower().endswith('.doc'):
        path = doc_to_docx(path)
    
    doc = Document(path)
    print("--- Paragraphs ---")
    for i, para in enumerate(doc.paragraphs):
        # We want to see EVERYTHING in paragraphs that looks like a placeholder
        if "(" in para.text or ")" in para.text:
            print(f"P{i}: {para.text}")
    
    print("\n--- Tables ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                if "(" in cell.text or ")" in cell.text:
                    print(f"  T{i} R{r} C{c}: {cell.text}")

if __name__ == "__main__":
    inspect_docx("1 Профинанс Д и Д ЕООД.doc")
